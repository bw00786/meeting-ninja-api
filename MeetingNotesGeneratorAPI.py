import os
import re
import logging
import requests
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from typing import Dict, Any, Optional

# Configure logging
logging.basicConfig(level=logging.INFO)

# Initialize Flask App
app = Flask(__name__)
CORS(app)

# Configuration from environment variables
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', '/path/to/upload/folder')
app.config['OUTPUT_FOLDER'] = os.environ.get('OUTPUT_FOLDER', '/path/to/output/folder')
PORT = int(os.environ.get('FLASK_PORT', 5000))  # Default to 5000 if not set

# LLM Service Configuration
PRIMARY_URL = os.environ.get('PRIMARY_LLM_URL')
BACKUP_URL = os.environ.get('BACKUP_LLM_URL')
PRIMARY_API_KEY = os.environ.get('PRIMARY_LLM_API_KEY')
BACKUP_API_KEY = os.environ.get('BACKUP_LLM_API_KEY')
PRIMARY_MODEL = os.environ.get('PRIMARY_LLM_MODEL')
BACKUP_MODEL = os.environ.get('BACKUP_LLM_MODEL')


# LLM Call Function for Meeting Minutes Generation
def generate_comprehensive_minutes(transcript: str) -> str:
    """
    Use LLM to generate comprehensive meeting minutes.
    
    Args:
        transcript (str): Full meeting transcript
    
    Returns:
        str: Formatted meeting minutes
    """
    try:
        # Payload for Anthropic's Claude 3 (adjust as needed for your specific LLM)
        payload = {
            "model": PRIMARY_MODEL,
            "max_tokens": 8000,
            "messages": [
                            {"role": "system", "content": "You are an assistant that helps analyze meeting transcripts."},
            {"role": "user", "content": f"Analyze the following meeting transcript:\n\n{transcript}\n\nProvide very detailed meeting minutes with details like the date of the meeting, the speakers, what were the highlights of what the speakers said, the category, any conclusions, next steps, and action items. Make sure you highlight Attendees, Categories, conclusions and Meeting Summary in bold. The headings should be supported by pdf format, place the bold text in Markdown format."}
                
            ]
        }
        
        # LLM Service Configuration (replace with your actual endpoint)
       
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {PRIMARY_API_KEY}"}
          
        
        
        # Add more detailed logging
        logging.info(f"Attempting to connect to {PRIMARY_URL}")
        logging.info(f"Using API Key: {PRIMARY_API_KEY[:5]}...")  # Partial key for security

        # Increase timeout and add retry mechanism
        response = requests.post(
            PRIMARY_URL, 
            json=payload, 
            headers=headers
           ## timeout=45,  # Increased timeout
          ##  verify=True  # Ensure SSL certificate verification
        )
        
        response.raise_for_status()  # This will raise an exception for bad HTTP status
        
        # Process successful response
        result = response.json()
         # Check for the structure and existence of 'choices'
        if 'choices' in result and len(result['choices']) > 0:
                return response.json()['choices'][0]['message']['content']
               
        else:
                # Log the specific error or issue with the response
                error_message = result.get('choices', [{'message': {'content': 'Unknown error'}}])[0]['message']['content']
                logging.error(f"LLM returned an error: {error_message}")
                return f"LLM error: {error_message}"
        return response.json()['choices'][0]['message']['content']

    except requests.exceptions.Timeout:
        logging.error("Connection to LLM service timed out")
        return "Error: Connection to LLM service timed out"
    
    except requests.exceptions.ConnectionError as e:
        logging.error(f"Network connectivity error: {e}")
        return f"Network Error: Unable to connect to LLM service. Details: {str(e)}"
    
    except requests.exceptions.RequestException as e:
        logging.error(f"Unexpected error in LLM service call: {e}")
        return f"Unexpected Error: {str(e)}"
    
    except Exception as e:
        logging.error(f"Unexpected general error: {e}")
        return f"General Error: {str(e)}"
class MeetingMinutesQA:
    def __init__(self, upload_folder, llm_urls, api_keys, models):
        self.upload_folder = upload_folder
        self.primary_url = llm_urls['primary']
        self.backup_url = llm_urls['backup']
        self.primary_api_key = api_keys['primary']
        self.backup_api_key = api_keys['backup']
        self.primary_model = models['primary']
        self.backup_model = models['backup']

    def _extract_text_from_file(self, filename: str) -> Optional[str]:
        """Extract text from a DOCX file."""
        try:
            full_path = os.path.join(self.upload_folder, filename)
            
            if filename.lower().endswith('.docx'):
                doc = Document(full_path)
                return '\n'.join([para.text for para in doc.paragraphs])
            
            else:
                raise ValueError(f"Unsupported file type: {filename}")
        
        except Exception as e:
            logging.error(f"Error extracting text: {e}")
            return None

    def _call_llm_service(self, transcript: str, question: str) -> Optional[str]:
        """Call LLM service with fallback mechanism."""
        # Primary LLM call
        try:
            payload = {
                "model": self.primary_model,
                "messages": [
                    {"role": "system", "content": "You are an AI assistant answering questions about meeting minutes."},
                    {"role": "user", "content": f"Meeting Minutes:\n{transcript}\n\nQuestion: {question}\n\nProvide a concise and direct answer based strictly on the meeting minutes."}
                ]
            }
            headers = {
                "Content-Type": "application/json",
                "Authorization": self.primary_api_key,
            }

            response = requests.post(self.primary_url, json=payload, headers=headers, timeout=30)
            response.raise_for_status()
            result = response.json()
            
            return result['choices'][0]['message']['content'] if 'choices' in result else None

        except Exception as primary_error:
            logging.warning(f"Primary LLM failed: {primary_error}")
            
            # Backup LLM call
            try:
                backup_payload = {
                    "prompt": f"Meeting Minutes:\n{transcript}\n\nQuestion: {question}\n\nAnswer:",
                    "max_tokens": 500,
                    "model": self.backup_model
                }
                backup_headers = {
                    "Content-Type": "application/json",
                    "Authorization": self.backup_api_key,
                }

                backup_response = requests.post(self.backup_url, json=backup_payload, headers=backup_headers, timeout=30)
                backup_response.raise_for_status()
                
                return backup_response.json().get('generated_text')

            except Exception as backup_error:
                logging.error(f"Backup LLM failed: {backup_error}")
                return None

def write_minutes_to_pdf(meeting_minutes: str, output_pdf_path: str) -> None:
    """Write the meeting minutes to a PDF file."""
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Set font for the body text
        pdf.set_font("Arial", size=12)

        # Add a title
        pdf.set_font("Arial", style='B', size=16)
        pdf.cell(200, 10, txt="Meeting Minutes", ln=True, align="C")
        pdf.ln(10)

        # Add the meeting minutes content
        pdf.set_font("Arial", size=14)
        for line in meeting_minutes.split('\n'):
            # Detect and process bold text in markdown-like format (**bold**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    pdf.set_font("Arial", style='B', size=14)
                    pdf.multi_cell(0, 10, part[2:-2])
                else:
                    pdf.set_font("Arial", size=14)
                    if part.strip():  # Only process non-empty parts
                        pdf.multi_cell(0, 10, part)

        pdf.output(output_pdf_path)
    except Exception as e:
        logging.error(f"Error writing PDF: {e}")
        raise

def write_minutes_to_docx(meeting_minutes: str, output_docx_path: str) -> None:
    """Write the meeting minutes to a DOCX file."""
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Meeting Minutes', 0)
        title.alignment = 1  # Center alignment

        # Process meeting minutes content
        for line in meeting_minutes.split('\n'):
            if line.startswith('**') and line.endswith('**'):
                # Bold paragraph
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line[2:-2])
                run.bold = True
            else:
                # Regular paragraph
                if line.strip():
                    doc.add_paragraph(line)

        # Style adjustments
        for paragraph in doc.paragraphs:
            paragraph.style.font.name = 'Calibri'
            paragraph.style.font.size = Pt(12)

        doc.save(output_docx_path)
    except Exception as e:
        logging.error(f"Error writing DOCX: {e}")
        raise

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(docx_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        raise

def analyze_meeting_transcript(transcript: str) -> str:
    """Analyze the meeting transcript (placeholder function)."""
    # This would typically involve an LLM call to summarize the transcript
    # For now, returning a simple placeholder
    return f"Meeting Transcript Summary:\n{transcript}"

# Initialize QA Handler
qa_handler = MeetingMinutesQA(
    upload_folder=app.config['UPLOAD_FOLDER'],
    llm_urls={'primary': PRIMARY_URL, 'backup': BACKUP_URL},
    api_keys={'primary': PRIMARY_API_KEY, 'backup': BACKUP_API_KEY},
    models={'primary': PRIMARY_MODEL, 'backup': BACKUP_MODEL}
)

@app.route('/generate_minutes', methods=['POST'])
def generate_meeting_minutes():
    """Generate meeting minutes from uploaded DOCX file with LLM fallback."""
    try:
        if 'docx_file' not in request.files:
            return jsonify({"error": "No DOCX file uploaded."}), 400

        file = request.files['docx_file']
        output_format = request.form.get('output_format', 'pdf')  # Default to PDF if not specified
        
        if file.filename == '' or not file.filename.endswith('.docx'):
            return jsonify({"error": "Invalid file. Please upload a DOCX file."}), 400

        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        base_filename = os.path.splitext(file.filename)[0]
        
        # Determine output file extension and path
        if output_format == 'pdf':
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{base_filename}.pdf")
            generate_func = write_minutes_to_pdf
        elif output_format == 'docx':
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{base_filename}_minutes.docx")
            generate_func = write_minutes_to_docx
        else:
            return jsonify({"error": "Invalid output format. Choose 'pdf' or 'docx'."}), 400

        file.save(docx_path)
        transcript = extract_text_from_docx(docx_path)
        meeting_minutes = generate_comprehensive_minutes(transcript)
        generate_func(meeting_minutes, output_path)

        return jsonify({
            "message": f"{output_format.upper()} generated successfully",
            "filename": os.path.basename(output_path),
            "docx_file": file.filename,
            "fullPath": output_path
        })

    except Exception as e:
        logging.error(f"Error in generate_meeting_minutes: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/ask_question', methods=['POST'])
def ask_question():
    """Enhanced question-answering endpoint."""
    try:
        data = request.get_json()
        question = data.get('question')
        filename = data.get('filename')

        if not question or not filename:
            return jsonify({"error": "Question and filename are required."}), 400

        # Extract text from the file
        transcript = qa_handler._extract_text_from_file(filename)
        
        if not transcript:
            return jsonify({"error": "Could not extract text from the file."}), 400

        # Get answer from LLM
        answer = qa_handler._call_llm_service(transcript, question)
        
        if not answer:
            return jsonify({"error": "Could not generate an answer."}), 500

        return jsonify({"answer": answer}), 200

    except Exception as e:
        logging.error(f"Error in ask_question: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/download_file/<filename>', methods=['GET'])
def download_file(filename: str):
    """Download generated PDF or DOCX file."""
    try:
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        if not os.path.exists(file_path):
            return jsonify({"error": f"File not found: {filename}"}), 404

        # Determine MIME type based on file extension
        if filename.lower().endswith('.pdf'):
            mimetype = 'application/pdf'
        elif filename.lower().endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return jsonify({"error": "Unsupported file type"}), 400

        return send_file(
            file_path,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logging.error(f"Error in download_file: {e}")
        return jsonify({"error": str(e)}), 500
@app.route('/')
def index():
    """
    Root route handler for the Meeting Minutes API.
    Provides basic information about the service.
    """
    return jsonify({
        "status": "API is running",
        "service": "Meeting Minutes Generation API",
        "version": "1.0.0",
        "endpoints": {
            "/generate_minutes": "POST - Generate meeting minutes from DOCX file",
            "/ask_question": "POST - Ask questions about a meeting transcript",
            "/download_file/<filename>": "GET - Download generated minutes"
        }
    }), 200
if __name__ == '__main__':
    # Use environment variable for port, with fallback to 5000
    app.run(host='0.0.0.0', port=PORT, debug=True)
